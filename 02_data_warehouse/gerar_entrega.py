#!/usr/bin/env python3
"""Gera arquitetura, capturas auditáveis, DOCX, HTML e PDF da entrega DW."""
from __future__ import annotations

import base64
import html
import json
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from playwright.sync_api import sync_playwright

ROOT = Path(__file__).resolve().parent
AZ = ROOT / "evidencias" / "azure"
CAP = ROOT / "evidencias" / "capturas_reais"
PORTAL = ROOT / "evidencias" / "capturas_portal_reais"
DOC_IMAGES = ROOT / "evidencias" / "capturas_documento"
FINAL = ROOT / "evidencias" / "entrega_final"
ICONS = ROOT / "evidencias" / "icones-azure-oficiais"
BRANDS = ROOT.parent / "01_cloud" / "evidencias" / "entrega_final" / "marcas"
CAP.mkdir(parents=True, exist_ok=True)
FINAL.mkdir(parents=True, exist_ok=True)
DOC_IMAGES.mkdir(parents=True, exist_ok=True)


def prepare_document_images() -> None:
    """Recorta somente a área vazia inferior; os originais permanecem intactos."""
    crop_heights = {
        "01_terminal_recursos_paas.png": 620,
        "02_terminal_blobs.png": 760,
        "05_terminal_dataflow_publicado.png": 1080,
        "06_terminal_derived_columns.png": 940,
        "09_terminal_pipeline_succeeded_2.png": 1240,
        "10_terminal_mysql_validacao.png": 1620,
        "11_terminal_txt_curated.png": 1980,
        "12_terminal_descarte_false.png": 300,
    }
    for name, height in crop_heights.items():
        subprocess.run([
            "magick", str(PORTAL / name), "-crop", f"2730x{height}+0+0", "+repage",
            str(DOC_IMAGES / name),
        ], check=True)


def load(name: str):
    return json.loads((AZ / name).read_text(encoding="utf-8"))


def text(name: str) -> str:
    return (AZ / name).read_text(encoding="utf-8")


def data_url(path: Path, mime: str = "image/svg+xml") -> str:
    return f"data:{mime};base64," + base64.b64encode(path.read_bytes()).decode()


ICON = {
    "rg": "10007-icon-service-Resource-Groups.svg",
    "adf": "10126-icon-service-Data-Factories.svg",
    "storage": "10086-icon-service-Storage-Accounts.svg",
    "mysql": "10122-icon-service-Azure-Database-MySQL-Server.svg",
    "monitor": "00001-icon-service-Monitor.svg",
}


def architecture_svg() -> str:
    def img(key: str, x: int, y: int, size: int = 58) -> str:
        return f'<image href="{data_url(ICONS / ICON[key])}" x="{x}" y="{y}" width="{size}" height="{size}"/>'
    return f'''<svg xmlns="http://www.w3.org/2000/svg" width="1600" height="900" viewBox="0 0 1600 900">
<defs><marker id="a" markerWidth="10" markerHeight="10" refX="9" refY="3" orient="auto"><path d="M0,0 L0,6 L9,3z" fill="#2563eb"/></marker><marker id="m" markerWidth="10" markerHeight="10" refX="9" refY="3" orient="auto"><path d="M0,0 L0,6 L9,3z" fill="#7c3aed"/></marker><style>.t{{font:700 29px Arial;fill:#102a43}}.h{{font:700 18px Arial;fill:#102a43}}.p{{font:14px Arial;fill:#334e68}}.b{{fill:#fff;stroke:#8db7df;stroke-width:2}}.e{{fill:none;stroke:#2563eb;stroke-width:3;marker-end:url(#a)}}.o{{fill:none;stroke:#7c3aed;stroke-width:3;stroke-dasharray:8 6;marker-end:url(#m)}}.l{{font:700 13px Arial;fill:#174ea6}}</style></defs>
<rect width="1600" height="900" fill="#f5f9fd"/><text x="800" y="44" text-anchor="middle" class="t">Predictfy — Arquitetura de ingestão executada</text>
<rect x="250" y="80" width="1310" height="750" rx="20" fill="#eef6ff" stroke="#0078d4" stroke-width="4"/>{img('rg',275,95,42)}<text x="325" y="123" class="h">Resource Group rg-predictfy-dw-sprint3-260823es</text>
<rect class="b" x="25" y="210" width="185" height="105" rx="12"/><text x="117" y="244" text-anchor="middle" class="h">CSV Predictfy</text><text x="117" y="270" text-anchor="middle" class="p">25.588 linhas</text><text x="117" y="293" text-anchor="middle" class="p">32 colunas</text>
<rect class="b" x="25" y="430" width="185" height="105" rx="12"/><text x="117" y="464" text-anchor="middle" class="h">Regras OLA</text><text x="117" y="490" text-anchor="middle" class="p">TXT pipe-delimited</text><text x="117" y="513" text-anchor="middle" class="p">P2 4h · P3 12h</text>
<rect class="b" x="320" y="285" width="260" height="180" rx="12"/>{img('storage',350,338)}<text x="430" y="327" class="h">Blob Storage</text><text x="430" y="355" class="p">container landing</text><text x="430" y="380" class="p">CSV + TXT de referência</text><text x="430" y="408" class="p">TLS 1.2 · acesso privado</text>
<rect class="b" x="675" y="205" width="330" height="330" rx="12"/>{img('adf',705,250)}<text x="790" y="244" class="h">Azure Data Factory</text><text x="790" y="274" class="p">Pipeline + Mapping Data Flow</text><text x="715" y="322" class="p">1 · projeção e limpeza</text><text x="715" y="350" class="p">2 · filtro de domínio</text><text x="715" y="378" class="p">3 · JOIN P2/P3</text><text x="715" y="406" class="p">4 · dez derivações</text><text x="715" y="434" class="p">5 · chave SHA-256</text><text x="715" y="462" class="p">6 · upsert + TXT</text><text x="715" y="500" class="p">Monitor: duas cargas Succeeded</text>
<rect class="b" x="1110" y="170" width="360" height="190" rx="12"/>{img('mysql',1140,225)}<text x="1225" y="213" class="h">Azure Database for MySQL</text><text x="1225" y="243" class="p">fato_incidente_predictfy</text><text x="1225" y="270" class="p">25.588 chaves únicas</text><text x="1225" y="297" class="p">upsert idempotente · TLS</text><text x="1225" y="324" class="p">Standard_B2s efêmero</text>
<rect class="b" x="1110" y="455" width="360" height="175" rx="12"/>{img('storage',1140,505)}<text x="1225" y="495" class="h">Blob Storage curated</text><text x="1225" y="525" class="p">incidentes_predictfy.txt</text><text x="1225" y="552" class="p">25.588 linhas · delimitador |</text><text x="1225" y="579" class="p">5.917.752 bytes</text>
<rect class="b" x="675" y="670" width="330" height="110" rx="12"/>{img('monitor',705,696,50)}<text x="775" y="705" class="h">ADF Monitor</text><text x="775" y="735" class="p">linhas, duração, partições</text><text x="775" y="760" class="p">execução e rastreabilidade</text>
<path class="e" d="M210 262 H270 V330 H320"/><path class="e" d="M210 482 H260 V420 H320"/><path class="e" d="M580 375 H675"/><path class="e" d="M1005 315 H1060 V265 H1110"/><path class="e" d="M1005 425 H1060 V542 H1110"/><path class="o" d="M840 535 V670"/><text x="595" y="358" class="l">leitura</text><text x="1018" y="250" class="l">upsert por ID</text><text x="1018" y="522" class="l">TXT curated</text>
<text x="800" y="866" text-anchor="middle" class="p">Fluxo real da Sprint 3: duas origens no landing → transformação governada no ADF → dois destinos persistentes na nuvem.</text></svg>'''


def dataflow_svg() -> str:
    nodes = [
        ("CSV", 35, 105, "sourceIncidentes"), ("Select", 300, 105, "selectContrato"),
        ("Filter", 565, 105, "filterValidos"), ("Derived", 830, 105, "derivePrioridade"),
        ("TXT", 35, 305, "sourceRegrasOla"), ("Select", 300, 305, "selectRegras"),
        ("Join", 830, 305, "joinRegrasOla"), ("Derived", 1100, 305, "deriveNegocio"),
        ("Select", 1100, 505, "selectDestino"), ("Alter Row", 830, 505, "alterRowUpsert"),
        ("MySQL", 565, 505, "sinkMySql"), ("TXT", 300, 505, "sinkTxt"),
    ]
    palette = {"CSV":"#0078d4", "TXT":"#0078d4", "Select":"#475569", "Filter":"#d97706", "Derived":"#7c3aed", "Join":"#0f766e", "Alter Row":"#be123c", "MySQL":"#087443"}
    boxes=[]
    for kind,x,y,name in nodes:
        color=palette.get(kind,"#087443")
        boxes.append(f'<rect x="{x}" y="{y}" width="205" height="82" rx="10" fill="#fff" stroke="{color}" stroke-width="3"/><text x="{x+102}" y="{y+31}" text-anchor="middle" class="k" fill="{color}">{kind}</text><text x="{x+102}" y="{y+59}" text-anchor="middle" class="n">{name}</text>')
    edges=[((240,146),(300,146)),((505,146),(565,146)),((770,146),(830,146)),
           ((240,346),(300,346)),((505,346),(830,346)),((932,187),(932,305)),
           ((1035,346),(1100,346)),((1202,387),(1202,505)),((1100,546),(1035,546)),
           ((830,546),(770,546))]
    lines=[]
    for (x1,y1),(x2,y2) in edges:
        mid=(x1+x2)//2
        lines.append(f'<path d="M{x1} {y1} H{mid} V{y2} H{x2}" fill="none" stroke="#2563eb" stroke-width="3" marker-end="url(#a)"/>')
    lines.append('<path d="M1202 587 V620 H402 V587" fill="none" stroke="#2563eb" stroke-width="3" marker-end="url(#a)"/>')
    return f'''<svg xmlns="http://www.w3.org/2000/svg" width="1360" height="680" viewBox="0 0 1360 680"><defs><marker id="a" markerWidth="10" markerHeight="10" refX="9" refY="3" orient="auto"><path d="M0,0 L0,6 L9,3z" fill="#2563eb"/></marker><style>.t{{font:700 26px Arial;fill:#102a43}}.k{{font:700 17px Arial}}.n{{font:15px Arial;fill:#334e68}}.lane{{font:700 14px Arial;fill:#64748b}}</style></defs><rect width="1360" height="680" fill="#f8fbfe"/><text x="680" y="48" text-anchor="middle" class="t">df_incidentes_predictfy — fluxo publicado no ADF</text><text x="35" y="88" class="lane">ORIGEM PRINCIPAL E LIMPEZA</text><text x="35" y="288" class="lane">REFERÊNCIA E INTEGRAÇÃO</text><text x="35" y="488" class="lane">PERSISTÊNCIA NOS DOIS DESTINOS</text>{''.join(lines)}{''.join(boxes)}<text x="680" y="650" text-anchor="middle" class="n">Representação fiel reconstruída do JSON coletado pela API; não é captura do ADF Studio.</text></svg>'''


def terminal_html(title: str, command: str, output: str) -> str:
    return f'''<!doctype html><html><head><meta charset="utf-8"><style>*{{box-sizing:border-box}}body{{margin:0;background:#07111f;color:#d7e7f4;font:15px/1.48 Menlo,monospace}}.bar{{height:52px;background:#075985;display:flex;align-items:center;padding:0 22px}}.title{{color:#fff;font:700 15px Arial}}pre{{margin:0;padding:22px 27px;white-space:pre-wrap;overflow-wrap:anywhere}}.prompt{{color:#67e8f9}}.cmd{{color:#fff;font-weight:700}}</style></head><body><div class="bar"><span class="title">QUADRO DE EVIDÊNCIA TÉCNICA · {html.escape(title)} · fonte Azure API/CLI</span></div><pre><span class="prompt">Comando consultado:</span> <span class="cmd">{html.escape(command)}</span>\n\n{html.escape(output)}</pre></body></html>'''


def summarize_dataflow() -> str:
    flow = load("07_dataflow_publicado.json")
    transforms = [item["name"] for item in flow.get("transformations", [])]
    lines = ["name: df_incidentes_predictfy", "type: MappingDataFlow", "", "sources:", "  - sourceIncidentes -> ds_csv_incidentes_landing", "  - sourceRegrasOla -> ds_txt_regras_ola", "", "transformations:"]
    lines.extend(f"  - {name}" for name in transforms)
    lines += ["", "sinks:", "  - sinkMySql -> fato_incidente_predictfy (upsert)", "  - sinkTxt -> curated/incidentes_predictfy.txt"]
    return "\n".join(lines)


def make_captures() -> None:
    (FINAL / "arquitetura_ingestao.svg").write_text(architecture_svg(), encoding="utf-8")
    (FINAL / "dataflow_canvas.svg").write_text(dataflow_svg(), encoding="utf-8")
    resources = text("01_recursos_paas.txt") + "\n\nDeployment Bicep: " + load("02_deployment_bicep.json")["state"]
    blobs = "LANDING\n" + json.dumps(load("04_blob_landing.json"), ensure_ascii=False, indent=2) + "\n\nCURATED\n" + json.dumps(load("05_blob_curated.json"), ensure_ascii=False, indent=2)
    runs = json.dumps(load("10_execucao_integral_1.json"), ensure_ascii=False, indent=2) + "\n\n" + json.dumps(load("11_execucao_integral_2.json"), ensure_ascii=False, indent=2)
    mysql = load("13_mysql_validacao.json")
    mysql_summary = json.dumps({k:mysql[k] for k in ("total_registros","chaves_unicas","prioridades_sem_regra","primeira_carga_apos_reexecucao","ultima_carga_apos_reexecucao","prioridades","ocorrencia_ola")}, ensure_ascii=False, indent=2)
    derived = (ROOT / "colunas_derivadas.txt").read_text(encoding="utf-8")
    curated = text("14_curated_sample.txt")
    cards = [
        ("01_recursos_paas.png", "Infraestrutura PaaS", "az resource list -g rg-predictfy-dw-sprint3-260823es", resources),
        ("02_blobs_landing_curated.png", "Azure Blob Storage", "az storage blob list --container landing/curated", blobs),
        ("04_colunas_derivadas.png", "Derived Column", "expressões publicadas no df_incidentes_predictfy", derived),
        ("05_pipeline_runs.png", "ADF Monitor", "pipeline runs + activity metrics", runs),
        ("06_mysql_validacao.png", "Azure Database for MySQL", "SELECT COUNT(*) / COUNT(DISTINCT ...) / GROUP BY", mysql_summary),
        ("07_txt_curated.png", "TXT curated", "download range incidentes_predictfy.txt", curated[:3500]),
        ("08_dataflow_api.png", "Data Flow publicado", "az datafactory data-flow show", summarize_dataflow()),
    ]
    with sync_playwright() as pw:
        browser = pw.chromium.launch()
        page = browser.new_page(viewport={"width": 1500, "height": 920}, device_scale_factor=1)
        page.set_default_timeout(120_000)
        for name,title,command,output in cards:
            page.set_content(terminal_html(title, command, output))
            page.screenshot(path=str(CAP / name), full_page=False)
        for source,name,size in ((FINAL/"arquitetura_ingestao.svg","03_arquitetura.png",(1600,900)),(FINAL/"dataflow_canvas.svg","03b_dataflow_canvas.png",(1360,680))):
            page.set_viewport_size({"width": size[0], "height": size[1]})
            page.goto(source.as_uri(), wait_until="load")
            page.screenshot(path=str(CAP/name), full_page=False)
        browser.close()


def figure(path: str, caption: str, folder: str = "capturas_reais") -> str:
    if (DOC_IMAGES / path).exists():
        folder = "capturas_documento"
    return f'<figure><img class="shot" src="../{folder}/{path}"><figcaption>{caption}</figcaption></figure>'


def command_note(command: str, purpose: str, requirement: str) -> str:
    return (f'<div class="command"><div><b>Comando executado</b></div><code>{html.escape(command)}</code>'
            f'<p><b>O que faz:</b> {purpose}<br><b>O que comprova:</b> {requirement}</p></div>')


def page(title: str, body: str) -> str:
    return f'<section class="page"><h1>{title}</h1>{body}</section>'


def report_html() -> str:
    r1, r2 = load("10_execucao_integral_1.json"), load("11_execucao_integral_2.json")
    destroyed = (AZ / "15_descarte_azure.txt").exists()
    destruction = html.escape(text("15_descarte_azure.txt")) if destroyed else "Descarte pendente até a revisão final."
    css='''@page{size:A4;margin:17mm 0 15mm}*{box-sizing:border-box}body{margin:0;background:#dce6ef;color:#102a43;font:13px/1.45 Arial,sans-serif}.page{width:210mm;margin:5mm auto;padding:8mm 17mm;background:#fff}.cover{height:297mm;margin:0 auto;padding:20mm;background:linear-gradient(145deg,#061a3a,#075985 62%,#00a4ef);color:#fff;display:flex;align-items:center;page-break-after:always}.cover h1{font-size:45px;color:#fff;border:0;padding:0}.cover h2{font-size:26px;color:#bde8ff}.cover p{font-size:18px;max-width:160mm}h1{font-size:24px;color:#082f57;border-left:5px solid #0078d4;padding-left:10px}h2{font-size:17px;color:#075985}h1,h2,table,pre,figure{break-inside:avoid}table{width:100%;border-collapse:collapse;margin:8px 0;font-size:11px}th{background:#075985;color:#fff}th,td{border:1px solid #bfd0df;padding:6px;text-align:left}pre{white-space:pre-wrap;background:#071525;color:#c8f7d4;border-radius:7px;padding:9px;font:9px/1.4 Menlo,monospace}.shot{display:block;max-width:100%;max-height:175mm;margin:auto;border:1px solid #c9d8e5;border-radius:5px}figcaption{font-size:10px;color:#627d98;margin-top:4px}.callout{padding:9px 11px;border-left:5px solid #0078d4;background:#eaf6ff}.ok{color:#087443;font-weight:700}@media print{body{background:#fff}.page{margin:0 auto}.cover{margin:-17mm auto -15mm}}'''
    css += '''.shot{border:0!important;border-radius:9px!important;box-shadow:none!important}figure{margin:7px 0 10px}h1{break-after:avoid!important;page-break-after:avoid!important}.command{break-inside:avoid;background:#f6f8fa;border:1px solid #d0d7de;border-left:4px solid #54aeff;border-radius:6px;padding:5px 8px;margin:5px 0;font-size:8px;line-height:1.3;color:#24292f}.command code{display:block;background:#eef1f4;color:#24292f;border:1px solid #d8dee4;border-radius:4px;padding:4px;margin-top:3px;white-space:pre-wrap}.command p{margin:3px 0 0}'''
    sections=[]
    sections.append('<section class="cover"><div><div>CHALLENGE LOCAWEB 2026 · SPRINT 3</div><h1>Predictfy</h1><h2>Data Warehousing &amp; Advanced Data Integration</h2><p>Azure Data Factory · Mapping Data Flow · MySQL · Blob Storage · ETL auditável</p></div></section>')
    sections.append(page("Identificação e objetivo",'''<p><b>Grupo:</b> Predictfy · <b>Turma:</b> 2TSCPW</p><table><tr><th>Integrante</th><th>RM</th></tr><tr><td>Elton Vinicios Almeida de Oliveira</td><td>562187</td></tr><tr><td>Emerson dos Santos Silva</td><td>562033</td></tr><tr><td>Kelvin Douglas Ribeiro Rabelo</td><td>561538</td></tr><tr><td>Pedro Henrique Simão Soares</td><td>562283</td></tr><tr><td>Vitor Lucas Mattos de Brito Mariano</td><td>562116</td></tr></table><p>Objetivo: integrar a matriz histórica de features de incidentes Predictfy com regras de OLA, transformar e enriquecer os registros no Azure Data Factory e persistir o mesmo contrato analítico em MySQL e TXT.</p>'''))
    sections.append(page("1. Arquitetura da solução de ingestão",figure("03_arquitetura.png","Figura 1 — arquitetura executada. CSV e regras TXT chegam ao landing; o ADF limpa, integra e deriva atributos; MySQL e TXT curated recebem 25.588 linhas. Ícones: Microsoft Azure Public Service Icons V24.")+'''<p>O Blob Storage funciona como Stage Area. O Mapping Data Flow é a única fronteira de transformação, evitando divergência lógica entre os dois sinks. O ADF Monitor registra status, duração, leitura e escrita.</p>'''))
    sections.append(page("2. Infraestrutura PaaS no Microsoft Azure",figure("01_terminal_recursos_paas.png","Figura 2 — captura nativa do Terminal com o inventário retornado pela Azure CLI.","capturas_portal_reais")+'''<table><tr><th>Serviço</th><th>Configuração executada</th></tr><tr><td>ADF e Storage</td><td>East US 2 · TLS 1.2 · identidade gerenciada</td></tr><tr><td>MySQL Flexible Server</td><td>Chile Central · MySQL 8.0.21 · Standard_B2s temporário</td></tr><tr><td>Containers</td><td>landing e curated · acesso público desabilitado</td></tr></table>'''+figure("02_terminal_blobs.png","Figura 3 — captura nativa do Terminal com a listagem real dos containers landing e curated.","capturas_portal_reais")))
    sections.append(page("3. Data Flow e processo ETL",figure("03_adf_studio_dataflow.png","Figura 4 — captura direta do ADF Studio. A bifurcação é precisa: selectDestino segue para Alter Row/MySQL e, diretamente, para TXT.","capturas_portal_reais")+figure("05_terminal_dataflow_publicado.png","Figura 5 — captura nativa do Terminal confirmando origens, transformações e destinos publicados.","capturas_portal_reais")+'''<p>O fluxo projeta o contrato necessário, valida hora/target/médias móveis, traduz prioridade, faz JOIN com duas regras OLA e cria atributos operacionais. Somente o ramo MySQL passa por Alter Row para o upsert por SHA-256; o ramo TXT parte diretamente de selectDestino.</p>'''))
    sections.append(page("4. Colunas derivadas",figure("06_terminal_derived_columns.png","Figura 6 — captura nativa do Terminal consultando a expressão deriveNegocio diretamente no Data Flow publicado.","capturas_portal_reais")+'''<table><tr><th>Derivação</th><th>Significado</th></tr><tr><td>incident_feature_id</td><td>chave determinística para idempotência</td></tr><tr><td>prioridade_codigo/desc</td><td>tradução do binário e enriquecimento pelo TXT</td></tr><tr><td>periodo_dia_desc</td><td>faixa operacional da hora</td></tr><tr><td>ocorrencia_ola</td><td>resultado histórico, não previsão</td></tr><tr><td>media_movel_delta/tendencia_volume</td><td>diferença 7d–30d e classificação</td></tr><tr><td>contexto_calendario</td><td>feriado, fim de semana ou dia útil</td></tr><tr><td>data_processamento/origem/arquivo</td><td>linhagem e auditoria</td></tr></table>'''))
    sections.append(page("5. Execução do pipeline e persistência",figure("08_adf_monitor_pipeline_runs.png",f"Figura 7 — captura direta do ADF Monitor com as duas execuções da carga.","capturas_portal_reais")+figure("09_terminal_pipeline_succeeded_2.png","Figura 8 — captura nativa do Terminal com status e métricas das duas execuções.","capturas_portal_reais")+figure("10_terminal_mysql_validacao.png","Figura 9 — consultas reais no MySQL após a reexecução idempotente.","capturas_portal_reais")+figure("11_terminal_txt_curated.png","Figura 10 — listagem e amostra real do TXT curated.","capturas_portal_reais")+'''<div class="callout"><b>Aceite:</b> 25.588 lidas; zero rejeitadas; 25.588 no MySQL; 25.588 no TXT; COUNT(*) = COUNT(DISTINCT incident_feature_id); segunda execução idempotente.</div>'''))
    sections.append(page("6. Comentários finais",'''<p>O desafio materializou um fluxo ETL completo e reproduzível. Os principais aprendizados foram parametrizar linked services com segurança, declarar projeções para o grafo Spark e distinguir funções do ADF das funções SQL. Para melhorar o desempenho do TXT de arquivo único, o upsert foi ajustado para batch de 10.000 e oito partições por hash. O resultado mantém rastreabilidade, separa landing de curated e comprova consistência entre MySQL e TXT.</p>'''))
    sections[3] = sections[3].replace('</section>',
        command_note('az resource list -g rg-predictfy-dw-sprint3-260823es --query "[].{Nome:name,Tipo:type,Regiao:location}" -o table',
                     'Lista os recursos efetivamente existentes no Resource Group.', 'Requisito 2 — infraestrutura PaaS.') +
        command_note('az storage blob list --account-name pfdw260823es -c landing/curated -o table',
                     'Lista os objetos persistidos nos containers de entrada e saída.', 'Requisitos 2 e 5 — Stage Area e TXT curated.') + '</section>')
    sections[4] = sections[4].replace('</section>',
        command_note('az datafactory data-flow show -g rg-predictfy-dw-sprint3-260823es --factory-name adf-predictfy-dw-260823es -n df_incidentes_predictfy',
                     'Consulta no Azure o Mapping Data Flow publicado e seus componentes.', 'Requisito 3 — Data Flow publicado e descrito.') + '</section>')
    sections[5] = sections[5].replace('</section>',
        command_note('az datafactory data-flow show ... | jq -r ".properties.scriptLines[] | select(endswith(\"~> deriveNegocio\"))"',
                     'Extrai diretamente do artefato publicado a expressão de colunas derivadas.', 'Requisito 4 — derivações executadas e documentadas.') + '</section>')
    sections[6] = sections[6].replace('</section>',
        command_note('az datafactory pipeline-run show ...; az datafactory activity-run query-by-pipeline-run ...',
                     'Consulta status, duração e métricas de escrita das duas execuções.', 'Requisito 5 — pipeline Succeeded e 25.588 linhas por destino.') +
        command_note('python3 scripts/collect_evidence.py; jq "{total_registros,chaves_unicas,prioridades_sem_regra}" evidencias/azure/13_mysql_validacao.json',
                     'Executa as consultas SQL sanitizadas e apresenta contagem, unicidade e cobertura das regras OLA.', 'Requisito 5 — persistência e idempotência no MySQL.') +
        command_note('az storage blob show ... incidentes_predictfy.txt; head -n 4 /tmp/predictfy_curated_sample.txt',
                     'Confirma tamanho e data do blob e mostra o cabeçalho e registros reais.', 'Requisito 5 — TXT persistido com colunas derivadas.') + '</section>')
    return '<!doctype html><html lang="pt-BR"><head><meta charset="utf-8"><title>Data Warehouse Predictfy</title><style>'+css+'</style></head><body>'+''.join(sections)+'</body></html>'


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instruction)
    run._r.append(fld)


def build_docx() -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21)
    sec.top_margin = sec.bottom_margin = Cm(2)
    sec.left_margin = sec.right_margin = Cm(2)
    sec.different_first_page_header_footer = True
    styles = doc.styles
    styles["Normal"].font.name = "Arial"; styles["Normal"].font.size = Pt(10)
    styles["Title"].font.name = "Arial"; styles["Title"].font.size = Pt(34); styles["Title"].font.color.rgb = RGBColor(0,120,212)
    styles["Heading 1"].font.name = "Arial"; styles["Heading 1"].font.size = Pt(20); styles["Heading 1"].font.color.rgb = RGBColor(8,47,87)
    header = sec.header.paragraphs[0]
    header.text = "Predictfy   |   DATA WAREHOUSING · SPRINT 3   |   FIAP"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Challenge Locaweb 2026   ·   Página "); add_field(footer,"PAGE"); footer.add_run(" de "); add_field(footer,"NUMPAGES")
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run("CHALLENGE LOCAWEB 2026 · SPRINT 3\n").bold=True
    p=doc.add_paragraph(style="Title"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run("Predictfy")
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Data Warehousing & Advanced Data Integration"); r.bold=True; r.font.size=Pt(20)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run("Azure Data Factory · Mapping Data Flow · MySQL · Blob Storage")
    doc.add_page_break()
    doc.add_heading("Identificação e objetivo", level=1)
    doc.add_paragraph("Grupo: Predictfy · Turma: 2TSCPW")
    table=doc.add_table(rows=1, cols=2); table.style="Table Grid"; table.rows[0].cells[0].text="Integrante"; table.rows[0].cells[1].text="RM"
    for name,rm in [("Elton Vinicios Almeida de Oliveira","562187"),("Emerson dos Santos Silva","562033"),("Kelvin Douglas Ribeiro Rabelo","561538"),("Pedro Henrique Simão Soares","562283"),("Vitor Lucas Mattos de Brito Mariano","562116")]:
        cells=table.add_row().cells; cells[0].text=name; cells[1].text=rm
    doc.add_paragraph("Integrar features históricas de incidentes Predictfy e regras de OLA, enriquecer no ADF e persistir o mesmo contrato analítico em MySQL e TXT.")
    content=[
        ("1. Arquitetura da solução de ingestão","03_arquitetura.png","CSV e regras TXT chegam ao landing; o ADF transforma e publica 25.588 linhas em MySQL e TXT curated.",CAP),
        ("2. Infraestrutura PaaS no Microsoft Azure","01_terminal_recursos_paas.png","Captura nativa do Terminal com ADF, Storage e MySQL.",PORTAL),
        ("2.1 Origens e destinos no Blob Storage","02_terminal_blobs.png","Captura nativa do Terminal: landing contém CSV/TXT e curated contém o TXT final.",PORTAL),
        ("3. Data Flow e processo ETL","03_adf_studio_dataflow.png","Captura direta do canvas no ADF Studio autenticado.",PORTAL),
        ("3.1 Publicação confirmada pela API","05_terminal_dataflow_publicado.png","Captura nativa do Terminal com origens, transformações e destinos.",PORTAL),
        ("4. Colunas derivadas","06_terminal_derived_columns.png","Expressão deriveNegocio consultada no Data Flow publicado.",PORTAL),
        ("5. Execução do pipeline","08_adf_monitor_pipeline_runs.png","Captura direta do ADF Monitor com as duas execuções.",PORTAL),
        ("5.1 Métricas da reexecução","09_terminal_pipeline_succeeded_2.png","Captura nativa do Terminal com status e contagens.",PORTAL),
        ("5.2 Persistência MySQL","10_terminal_mysql_validacao.png","25.588 registros, 25.588 chaves únicas e zero prioridades sem regra.",PORTAL),
        ("5.3 Persistência TXT","11_terminal_txt_curated.png","Listagem e amostra real do TXT pipe-delimited.",PORTAL),
    ]
    explanations={
        "1. Arquitetura da solução de ingestão":"O Blob funciona como Stage Area; uma única transformação alimenta os dois destinos e o ADF Monitor conserva rastreabilidade.",
        "2. Infraestrutura PaaS no Microsoft Azure":"ADF e Storage foram provisionados em East US 2; MySQL 8.0.21 em Chile Central, com TLS e credenciais somente no ambiente.",
        "3. Data Flow e processo ETL":"O fluxo projeta, valida, integra P2/P3 e deriva dez atributos. O MySQL recebe upsert por SHA-256; o TXT recebe o mesmo stream.",
        "4. Colunas derivadas":"As derivações traduzem atributos técnicos para conceitos operacionais, preservam linhagem e deixam explícito que ocorrência OLA é histórica. As expressões completas também estão entregues em colunas_derivadas.txt.",
        "5. Execução do pipeline":"A segunda carga atualizou as mesmas chaves, comprovando idempotência sem duplicação.",
    }
    commands={
        "2. Infraestrutura PaaS no Microsoft Azure":("az resource list -g rg-predictfy-dw-sprint3-260823es -o table","Lista os recursos PaaS reais do Resource Group e comprova o requisito 2."),
        "2.1 Origens e destinos no Blob Storage":("az storage blob list --account-name pfdw260823es -c landing/curated -o table","Lista os arquivos de entrada e saída e comprova Stage Area e persistência."),
        "3. Data Flow e processo ETL":("az datafactory data-flow show -g rg-predictfy-dw-sprint3-260823es --factory-name adf-predictfy-dw-260823es -n df_incidentes_predictfy","Consulta o Data Flow publicado e comprova o requisito 3."),
        "4. Colunas derivadas":("az datafactory data-flow show ... | jq '.properties.scriptLines[] | select(endswith(\"~> deriveNegocio\"))'","Extrai as derivações do artefato publicado e comprova o requisito 4."),
        "5.1 Métricas da reexecução":("az datafactory pipeline-run show ...; az datafactory activity-run query-by-pipeline-run ...","Consulta status, duração e linhas escritas nas duas execuções."),
        "5.2 Persistência MySQL":("python3 scripts/collect_evidence.py; jq evidencias/azure/13_mysql_validacao.json","Executa validações SQL sanitizadas e comprova contagem, unicidade e regras OLA."),
        "5.3 Persistência TXT":("az storage blob show ...; head -n 4 /tmp/predictfy_curated_sample.txt","Confirma o blob curated e apresenta uma amostra real do conteúdo."),
    }
    for title,image_name,caption,image_dir in content:
        doc.add_heading(title, level=1 if ".1 " not in title and ".2 " not in title else 2)
        if title in explanations: doc.add_paragraph(explanations[title])
        prepared = DOC_IMAGES / image_name
        doc.add_picture(str(prepared if prepared.exists() else image_dir/image_name), width=Inches(6.4))
        p=doc.add_paragraph("Figura — "+caption); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].italic=True; p.runs[0].font.size=Pt(8)
        if title in commands:
            command, purpose = commands[title]
            p=doc.add_paragraph(); p.add_run("Comando executado: ").bold=True; p.add_run(command).font.name="Consolas"
            p=doc.add_paragraph(); p.add_run("O que faz e comprova: ").bold=True; p.add_run(purpose)
    doc.add_page_break()
    doc.add_heading("6. Comentários finais",level=1)
    doc.add_paragraph("O desafio consolidou o uso do Azure Data Factory para ETL real. Foi necessário corrigir projeções do grafo Spark, validar funções do Mapping Data Flow e otimizar o upsert com batch de 10.000 e oito partições por hash. O resultado separa landing/curated, preserva linhagem e mantém os dois destinos consistentes.")
    doc.add_paragraph("Aceite: 25.588 lidas; zero rejeitadas; 25.588 no MySQL; 25.588 no TXT; chaves únicas; duas execuções Succeeded.")
    doc.save(ROOT / "DataWarehouse_Locaweb_Sprint3_Predictfy.docx")


def build_pdf() -> None:
    html_path=FINAL/"RELATORIO_ENTREGA_COMPLETO.html"
    html_path.write_text(report_html(),encoding="utf-8")
    predictfy=data_url(BRANDS/"predictfy-logo.svg")
    fiap=data_url(BRANDS/"fiap-oficial.ico","image/x-icon")
    header=f'<div style="width:100%;padding:0 12mm;font:9px Arial;color:#12324f"><div style="display:flex;justify-content:space-between;border-bottom:1px solid #bdd3e5;padding:2mm"><span><img src="{predictfy}" style="width:16px;vertical-align:middle"> Predictfy</span><span>DATA WAREHOUSING · SPRINT 3</span><span style="color:#ed145b"><img src="{fiap}" style="width:15px;vertical-align:middle"> FIAP</span></div></div>'
    footer='<div style="width:100%;padding:0 12mm;font:8px Arial;color:#526d82"><div style="display:flex;justify-content:space-between;border-top:1px solid #bdd3e5;padding:2mm"><span>Challenge Locaweb 2026</span><span>Data Warehousing &amp; Advanced Data Integration</span><span>Página <span class="pageNumber"></span> de <span class="totalPages"></span></span></div></div>'
    with tempfile.TemporaryDirectory(prefix="predictfy_dw_pdf_") as tmp:
        cover_pdf=Path(tmp)/"cover.pdf"; body_pdf=Path(tmp)/"body.pdf"
        with sync_playwright() as pw:
            browser=pw.chromium.launch()
            cover=browser.new_page(); cover.goto(html_path.as_uri(),wait_until="networkidle"); cover.add_style_tag(content="@page{margin:0!important}.page:not(.cover){display:none!important}.cover{margin:0!important}"); cover.pdf(path=str(cover_pdf),format="A4",print_background=True,display_header_footer=False,margin={"top":"0","right":"0","bottom":"0","left":"0"})
            body=browser.new_page(); body.goto(html_path.as_uri(),wait_until="networkidle"); body.add_style_tag(content="@page{size:A4 portrait!important}.cover{display:none!important}"); body.pdf(path=str(body_pdf),format="A4",print_background=True,display_header_footer=True,header_template=header,footer_template=footer,margin={"top":"17mm","right":"0","bottom":"15mm","left":"0"})
            browser.close()
        subprocess.run(["pdfunite",str(cover_pdf),str(body_pdf),str(ROOT/"DataWarehouse_Locaweb_Sprint3_Predictfy.pdf")],check=True)


def main() -> None:
    prepare_document_images(); make_captures(); build_docx(); build_pdf()
    print("DOCX, PDF, arquitetura e capturas gerados.")

if __name__ == "__main__": main()
