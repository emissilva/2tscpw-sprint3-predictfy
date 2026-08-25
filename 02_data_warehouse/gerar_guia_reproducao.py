#!/usr/bin/env python3
"""Converte o guia Markdown em DOCX e PDF com o padrão visual Predictfy."""
from pathlib import Path
import html
import re
from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from playwright.sync_api import sync_playwright

ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "Passo_a_Passo_Reproducao_DataWarehouse_Predictfy.md"
CANONICAL = ROOT / "PASSO_A_PASSO_EXECUCAO.md"
DOCX = ROOT / "Passo_a_Passo_Reproducao_DataWarehouse_Predictfy.docx"
PDF = ROOT / "Passo_a_Passo_Reproducao_DataWarehouse_Predictfy.pdf"


def blocks(text: str):
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            language = line[3:].strip()
            i += 1
            code = []
            while i < len(lines) and not lines[i].startswith("```"):
                code.append(lines[i]); i += 1
            yield "code", language, "\n".join(code)
        elif line.startswith("#"):
            match = re.match(r"^(#+)\s+(.*)$", line)
            yield "heading", len(match.group(1)), match.group(2)
        elif re.match(r"^\d+\.\s+", line):
            yield "number", 0, re.sub(r"^\d+\.\s+", "", line)
        elif line.startswith("> "):
            yield "quote", 0, line[2:]
        elif line.strip():
            paragraph = [line.strip()]
            while i + 1 < len(lines) and lines[i + 1].strip() and not re.match(r"^(#|```|> |\d+\.\s+)", lines[i + 1]):
                i += 1; paragraph.append(lines[i].strip())
            yield "paragraph", 0, " ".join(paragraph)
        i += 1


def plain(value: str) -> str:
    return re.sub(r"`([^`]+)`", r"\1", value).replace("**", "")


def make_docx():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(.65); sec.bottom_margin = Inches(.65)
    sec.left_margin = Inches(.75); sec.right_margin = Inches(.75)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"; normal.font.size = Pt(9.5); normal.font.color.rgb = RGBColor(36, 41, 47)
    for name, size, color in (("Title", 27, (8,47,87)), ("Heading 1", 18, (8,47,87)), ("Heading 2", 13, (7,89,133))):
        style = doc.styles[name]; style.font.name = "Arial"; style.font.size = Pt(size); style.font.color.rgb = RGBColor(*color)
    header = sec.header.paragraphs[0]
    header.text = "Predictfy   |   GUIA DE REPRODUÇÃO MANUAL   |   FIAP"
    header.alignment = 1
    footer = sec.footer.paragraphs[0]
    footer.text = "Data Warehousing · Sprint 3"
    footer.alignment = 1
    for kind, level, value in blocks(SOURCE.read_text(encoding="utf-8")):
        value = plain(value)
        if kind == "heading":
            doc.add_heading(value, 0 if level == 1 else min(level - 1, 2))
        elif kind == "code":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
            shade = OxmlElement("w:shd"); shade.set(qn("w:fill"), "F6F8FA"); p._p.get_or_add_pPr().append(shade)
            border = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
            for key, val in (("val","single"),("sz","4"),("space","2"),("color","D0D7DE")): bottom.set(qn(f"w:{key}"), val)
            border.append(bottom); p._p.get_or_add_pPr().append(border)
            run = p.add_run(value); run.font.name = "Consolas"; run.font.size = Pt(8); run.font.color.rgb = RGBColor(36,41,47)
        elif kind == "number":
            doc.add_paragraph(value, style="List Number")
        elif kind == "quote":
            p = doc.add_paragraph(value); p.style = doc.styles["Quote"]
        else:
            doc.add_paragraph(value)
    doc.save(DOCX)


def inline(value: str) -> str:
    safe = html.escape(value)
    return re.sub(r"`([^`]+)`", r"<code>\1</code>", safe)


def make_pdf():
    parts = []
    list_open = False
    for kind, level, value in blocks(SOURCE.read_text(encoding="utf-8")):
        if kind != "number" and list_open:
            parts.append("</ol>"); list_open = False
        if kind == "heading":
            tag = "h1" if level == 1 else "h2" if level == 2 else "h3"
            parts.append(f"<{tag}>{inline(value)}</{tag}>")
        elif kind == "code":
            parts.append(f"<pre><code>{html.escape(value)}</code></pre>")
        elif kind == "number":
            if not list_open: parts.append("<ol>"); list_open = True
            parts.append(f"<li>{inline(value)}</li>")
        elif kind == "quote":
            parts.append(f"<blockquote>{inline(value)}</blockquote>")
        else:
            parts.append(f"<p>{inline(value)}</p>")
    if list_open: parts.append("</ol>")
    css = """@page{size:A4;margin:18mm 17mm 17mm}*{box-sizing:border-box}body{font:10px/1.45 Arial;color:#24292f;margin:0}h1{font-size:25px;color:#082f57;margin:0 0 14px}h2{font-size:18px;color:#082f57;border-left:4px solid #54aeff;padding-left:8px;margin:19px 0 7px;break-after:avoid}h3{font-size:13px;color:#075985;margin:13px 0 5px;break-after:avoid}p{margin:5px 0}pre{background:#f6f8fa;color:#24292f;border:1px solid #d0d7de;border-radius:7px;padding:7px 9px;font:8px/1.35 Menlo,monospace;white-space:pre-wrap;overflow-wrap:anywhere;break-inside:avoid;margin:5px 0}p code,li code,blockquote code{background:#eef1f4;border-radius:3px;padding:1px 3px}blockquote{background:#eaf6ff;border-left:4px solid #54aeff;margin:7px 0;padding:7px 10px}ol{margin:5px 0 8px;padding-left:22px}li{margin:3px 0}"""
    document = f"<!doctype html><html lang='pt-BR'><head><meta charset='utf-8'><style>{css}</style></head><body>{''.join(parts)}</body></html>"
    with sync_playwright() as pw:
        browser = pw.chromium.launch()
        page = browser.new_page()
        page.set_content(document, wait_until="load")
        page.pdf(path=str(PDF), format="A4", print_background=True, display_header_footer=True,
                 header_template="<div style='width:100%;padding:0 17mm;font:8px Arial;color:#526d82'>Predictfy · Guia de reprodução manual</div>",
                 footer_template="<div style='width:100%;padding:0 17mm;text-align:right;font:8px Arial;color:#526d82'>Página <span class='pageNumber'></span> de <span class='totalPages'></span></div>",
                 margin={"top":"18mm","right":"0","bottom":"17mm","left":"0"})
        browser.close()


if __name__ == "__main__":
    CANONICAL.write_text(SOURCE.read_text(encoding="utf-8"), encoding="utf-8")
    make_docx()
    make_pdf()
    print(f"Gerados: {DOCX.name} e {PDF.name}")
